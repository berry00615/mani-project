#!/usr/bin/env python3
"""Build the final Chinese research report for TwoRobotPickCube-v1."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/two_robot_pick_cube_final_report_zh.docx")
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 99, 110)
LIGHT = "F2F4F7"
CONTENT_DXA = 9360
TABLE_INDENT = 120


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def fixed_table(doc, headers, rows, widths):
    assert sum(widths) == CONTENT_DXA
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for j, text in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_width(cell, widths[j])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), LIGHT)
        cell._tc.get_or_add_tcPr().append(shade)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(text), 10, True, DARK)
    for row in rows:
        cells = table.add_row().cells
        for j, text in enumerate(row):
            set_cell_width(cells[j], widths[j])
            set_cell_margins(cells[j])
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(str(text)), 9.5)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def para(doc, text, bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_label and text.startswith(bold_label):
        set_font(p.add_run(bold_label), 11, True)
        set_font(p.add_run(text[len(bold_label):]), 11)
    else:
        set_font(p.add_run(text), 11)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt({1: 16, 2: 12, 3: 8}[level])
    p.paragraph_format.space_after = Pt({1: 8, 2: 6, 3: 4}[level])
    run = p.add_run(text)
    set_font(run, {1: 16, 2: 13, 3: 12}[level], True, BLUE if level < 3 else DARK)
    return p


def page_break(doc):
    doc.add_page_break()


def build():
    if OUT.exists():
        raise FileExistsError(OUT)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, before, after, color in (
        (1, 16, 16, 8, BLUE), (2, 13, 12, 6, BLUE), (3, 12, 8, 4, DARK)
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("MANI-PROJECT  ·  TASK 6  ·  FINAL REPORT"), 8.5, True, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("TwoRobotPickCube-v1 · ManiSkill 3.0.1 · 2026-07-24"), 8.5, False, MUTED)

    # Editorial cover pattern with restrained technical-report styling.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(100)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("MANI-PROJECT · 任务六"), 11, True, BLUE)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("TwoRobotPickCube-v1"), 28, True, DARK)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("双机械臂协作搬运完整研究报告"), 17, True, BLUE)
    p.paragraph_format.space_after = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("PPO-fast · 多路线对照 · 全 checkpoint 筛选 · 1,000 环境正式评估"), 11, False, MUTED)
    p.paragraph_format.space_after = Pt(88)
    fixed_table(doc, ["项目", "最终结论"], [
        ("环境", "ManiSkill 3.0.1 · TwoRobotPickCube-v1"),
        ("最佳方法", "TRPC-O5 · O3 低学习率精修"),
        ("严格成功", "1,000 / 1,000"),
        ("模型哈希", "ebe5d686…1aa63c"),
    ], [2600, 6760])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("2026 年 7 月 24 日"), 10.5, True, MUTED)

    page_break(doc)
    heading(doc, "执行摘要", 1)
    para(doc, "本任务要求两台 Panda 机械臂完成空间受限的协作搬运：红色方块只在左臂工作区内初始化，而随机目标只在右臂工作区内可达。策略必须学会左臂交接、右臂抓取、目标运输与最终静止的完整时序。")
    para(doc, "研究从官方 PPO-fast 基线出发，系统比较三种折扣设置，并对最佳从零训练模型进行两档低学习率精修。最终 TRPC-O5 final 在选择集和独立集均达到 100/100，在 matched seed 20260723 的 1,000 环境正式评估中达到 1,000/1,000，且所有阶段指标均为 100%。")
    fixed_table(doc, ["结论项", "结果"], [
        ("最佳模型", "TRPC-O5 final"),
        ("正式严格成功", "1,000 / 1,000（100%）"),
        ("平均最小目标距离", "2.827 mm"),
        ("平均 normalized-dense return", "83.308"),
        ("失败分类", "全部为 0"),
        ("模型 SHA-256", "ebe5d6865c4d7282f428931258144ac73ec37add7b94c9578de3e8448d1aa63c"),
    ], [2700, 6660])
    heading(doc, "关键研究发现", 2)
    para(doc, "第一，官方 gamma=0.8 对需要多阶段信用分配的双臂交接任务偏短。第二，gamma=0.95 比 gamma=0.99 更稳定地学习完整协作链。第三，成功率接近饱和后，较温和的 5e-5 精修优于 1e-4；后者虽然回报提高，却引入了新的掉落和未抓取失败。")

    page_break(doc)
    heading(doc, "一、任务与严格成功定义", 1)
    para(doc, "环境为 ManiSkill 3.0.1 原生 TwoRobotPickCube-v1，最大 episode 长度 100。两台 panda_wristcam 以 pd_joint_delta_pos 控制，原始多智能体 Dict 动作经官方 FlattenActionSpaceWrapper 合并为 16 维联合动作；state 观测为 66 维。")
    fixed_table(doc, ["项目", "定义"], [
        ("随机化", "方块位于左侧工作区；目标位于右侧工作区且高度随机"),
        ("放置条件", "方块中心到目标中心距离 ≤ 0.025 m"),
        ("静止条件", "右臂满足原生 is_static(0.2)"),
        ("严格成功", "放置条件与静止条件同时成立"),
        ("奖励", "五阶段 dense reward，成功原始值 21，normalized 后除以 21"),
    ], [2300, 7060])
    heading(doc, "五阶段奖励结构", 2)
    para(doc, "阶段 1：左臂接近并将方块送过交接线；阶段 2：右臂接近并抓取，同时左臂避让；阶段 3：右臂运输，左臂回位；阶段 4：进入目标邻域；阶段 5：精确放置并保持静止。")
    heading(doc, "固定评估协议", 2)
    para(doc, "所有 checkpoint 先在 seed 0 的 100 个并行环境中全量筛选，再在 seed 10000 的独立 100 环境中复验；最终候选在 seed 20260723 的 1,000 环境上正式比较。在线评估只作健康指标，final 不自动晋升。")

    page_break(doc)
    heading(doc, "二、实验设计与工程保障", 1)
    fixed_table(doc, ["设置", "值"], [
        ("训练器", "ManiSkill 官方 PPO-fast"),
        ("并行环境", "1,024"),
        ("rollout", "100"),
        ("PPO epochs / minibatches", "8 / 32"),
        ("actor / critic", "三层 256 单元 Tanh MLP"),
        ("CUDA graphs", "启用"),
        ("基线总步数", "50M"),
        ("精修步数", "追加 10M"),
    ], [3000, 6360])
    heading(doc, "持久化与证据链", 2)
    para(doc, "所有长任务运行于服务器 tmux；SSH 断开和本地 Mac 合盖不影响训练。每个实验保留独立配置、完整命令、PID、GPU 映射、stdout/stderr、TensorBoard、三分钟 GPU 监控、全部周期 checkpoint、final、评估 JSON/CSV 和失败分类。")
    heading(doc, "训练前基线", 2)
    fixed_table(doc, ["策略", "严格成功", "越过交接线", "平均回报", "失败"], [
        ("随机动作", "0/100", "0%", "1.165", "100 次未交接"),
        ("未训练官方 actor", "0/100", "0%", "1.298", "100 次未交接"),
    ], [2600, 1600, 1800, 1500, 1860])
    para(doc, "两次启动负结果也被保留：安全缓存目录中的 PhysX 下载卡住；原始 Dict 动作空间无法直接读取 shape。后者通过采用官方 flatten wrapper 修复。")

    page_break(doc)
    heading(doc, "三、从零训练方法比较", 1)
    fixed_table(doc, ["方法", "gamma / GAE", "步骤", "选择集", "独立集", "正式评估"], [
        ("O1 官方", "0.8 / 0.9", "50M", "91/100", "—", "—"),
        ("O2 长信用", "0.99 / 0.95", "50M", "100/100", "99/100", "992/1000"),
        ("O3 中等信用", "0.95 / 0.95", "50M", "100/100", "100/100", "998/1000"),
    ], [1900, 1700, 1100, 1500, 1500, 1660])
    heading(doc, "O2 失败画像", 2)
    para(doc, "O2 ckpt_476 的 8 次失败包括 1 次未抓取和 7 次到位未静止。它已经学会交接和运输，但 gamma=0.99 下末端动作略显拖沓，静止条件成为主要瓶颈。")
    heading(doc, "O3 改进", 2)
    para(doc, "O3 ckpt_476 将正式成功率提升至 99.8%，仅余 1 次未抓取和 1 次到位未静止。其两组 100 seeds 均为满分，说明提升不是选择集偶然峰值。")
    fixed_table(doc, ["候选", "严格成功", "未抓取", "到位未静止", "平均回报"], [
        ("O2 ckpt_476", "992/1000", "1", "7", "78.715"),
        ("O3 ckpt_476", "998/1000", "1", "1", "81.796"),
    ], [2500, 1900, 1500, 1900, 1560])

    page_break(doc)
    heading(doc, "四、低学习率精修与最终晋升", 1)
    para(doc, "O4 与 O5 均从 O3 ckpt_476 的不可变副本启动，重置优化器，追加 10M 步并线性退火。实验只改变初始学习率，以隔离精修强度的影响。")
    fixed_table(doc, ["方法", "学习率", "两组 100 seeds", "正式成功", "失败", "结论"], [
        ("O4", "1e-4", "100 + 100", "997/1000", "1 掉落 + 2 未抓取", "回归，不晋升"),
        ("O5", "5e-5", "100 + 100", "1000/1000", "0", "最终晋升"),
    ], [1100, 1200, 1800, 1700, 2100, 1460])
    heading(doc, "为何不是按回报选模", 2)
    para(doc, "O4 的平均回报和目标精度均有改善，但严格成功反而从 998/1000 降到 997/1000。这证明 shaped reward 不能替代原生 success。O5 只有在 matched-seed 正式比较达到 1000/1000 后才被复制为 best。")
    heading(doc, "最终阶段指标", 2)
    fixed_table(doc, ["指标", "O5 final"], [
        ("曾越过交接线", "100.0%"),
        ("曾被右臂抓取", "100.0%"),
        ("曾进入目标范围", "100.0%"),
        ("到位且静止", "100.0%"),
        ("平均最小目标距离", "2.827 mm"),
        ("平均回报", "83.308"),
        ("互斥失败总数", "0"),
    ], [4400, 4960])

    page_break(doc)
    heading(doc, "五、模型、视频与复现信息", 1)
    heading(doc, "晋升模型", 2)
    para(doc, "本地：checkpoints/ppo_two_robot_pick_cube_o5_lr5e5_10m/best_o5_final_1000of1000_matchedseed.pt")
    para(doc, "服务器：runs/two_robot_pick_cube_TRPC-O5_o3ckpt476_lr5e5_10m_seed0_20260723/best_o5_final_1000of1000_matchedseed.pt")
    para(doc, "SHA-256：ebe5d6865c4d7282f428931258144ac73ec37add7b94c9578de3e8448d1aa63c")
    heading(doc, "视频证据", 2)
    para(doc, "seeds 0–19 共 20 个确定性 episode 均通过 strict success。每个 MP4 已验证为 H.264、512×512、30 fps、101 帧、3.3667 秒。目录：videos/two_robot_pick_cube/o5_best_20_successes_20260724。")
    heading(doc, "复现实验入口", 2)
    para(doc, "训练器：third_party/ManiSkill-v3.0.1-official/ppo/ppo_fast.py。评估器：scripts/evaluate_official_two_robot_pick_cube.py。视频脚本：scripts/record_two_robot_pick_cube_best.py。完整运行记录：docs/two_robot_pick_cube_experiment_log.md。")
    heading(doc, "结论", 1)
    para(doc, "任务六已完成。最佳策略在固定选择集、独立复验集和 matched-seed 1,000 环境正式评估中均无失败。研究结果支持在长时序双臂协作中采用中等折扣，并在接近饱和后使用非常低的学习率进行保守精修。")

    doc.core_properties.title = "TwoRobotPickCube-v1 双机械臂协作搬运完整研究报告"
    doc.core_properties.subject = "ManiSkill PPO research report"
    doc.core_properties.author = "mani-project"
    doc.core_properties.keywords = "ManiSkill, PPO, TwoRobotPickCube, multi-robot collaboration"
    doc.core_properties.comments = "Generated from retained experiment artifacts."
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
