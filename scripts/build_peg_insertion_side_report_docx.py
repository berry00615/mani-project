#!/usr/bin/env python3
"""Build the Chinese final report for PegInsertionSide-v1 without touching prior reports."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT / "docs/mani_project_full_research_report_stage6_9.docx"
OUTPUT = ROOT / "docs/peg_insertion_side_final_report_zh.docx"

BLUE = "365F91"
MID_BLUE = "4F81BD"
LIGHT_BLUE = "DCE6F1"
PALE_BLUE = "EEF3F8"
PALE_GREEN = "E2F0D9"
PALE_RED = "FCE4D6"
GRAY = "666666"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def set_font(run, name="Arial", size=10.5, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", bold_prefix=None, align=None, space_after=5, indent=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        for r in p.runs:
            set_font(r, size=10)
        if not p.runs:
            set_font(p.add_run(text), size=10)
        else:
            p.runs[0].text = text


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_table(doc, headers, rows, widths, highlight_rows=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = str(value)
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_font(r, size=font_size, bold=True, color=WHITE)
    highlight_rows = highlight_rows or {}
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        fill = highlight_rows.get(ridx, PALE_BLUE if ridx % 2 else WHITE)
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cells[i], fill)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    set_font(r, size=font_size)
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(f"{label}｜")
    set_font(r, bold=True, color=BLUE)
    r = p.add_run(text)
    set_font(r, size=10)
    set_table_width(table, [9220])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_font(run, size=8, color=GRAY)


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.70)
    sec.bottom_margin = Inches(0.70)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    sec.header_distance = Inches(0.30)
    sec.footer_distance = Inches(0.30)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 17, BLUE, 13, 5),
        ("Heading 2", 14, MID_BLUE, 8, 3),
        ("Heading 3", 12, MID_BLUE, 6, 2),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(10)
    header = sec.header
    p = header.paragraphs[0]
    p.text = "MANI-PROJECT  ·  TASK 3  ·  PEG INSERTION SIDE"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        set_font(r, size=7.5, bold=True, color=GRAY)
    footer = sec.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def build():
    doc = Document(REFERENCE)
    clear_body(doc)
    configure_styles(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于 ManiSkill 官方 PPO-fast 的\n侧向插销装配策略")
    set_font(r, size=25, bold=True, color="17365D")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("PegInsertionSide-v1 任务三完整研究报告")
    set_font(r, size=16, bold=True, color=MID_BLUE)
    add_callout(
        doc,
        "最终结论",
        "从零训练的官方 75M 步 O1 最佳模型在 matched-seed 1,000 环境中严格成功 175/1,000；"
        "以该候选继续进行 25M 步低学习率 O2 训练后，最佳 checkpoint 达到 991/1,000（99.1%），"
        "绝对提升 81.6 个百分点。",
        PALE_GREEN,
    )
    meta = [
        ("项目", "mani-project"),
        ("任务", "ManiSkill 3.0.1 · PegInsertionSide-v1"),
        ("官方参考", "commit a4a4f9272ad64b1564035874b605ceb687b63ed8"),
        ("作者", "黄泊睿"),
        ("报告日期", "2026 年 7 月 23 日"),
        ("最佳模型", "O2 ckpt_551 · 991/1,000 strict success"),
    ]
    add_table(doc, ["项目字段", "内容"], meta, [2100, 7120], font_size=9.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("任务三 · 冻结版研究与交付报告")
    set_font(r, size=10, bold=True, color=GRAY)
    doc.add_page_break()

    add_heading(doc, "摘要", 1)
    add_para(
        doc,
        "本报告记录 mani-project 第三阶段 PegInsertionSide-v1 的完整研究过程。该任务要求 Panda "
        "机械臂抓取横置插销，完成轴向姿态与孔口横向对准，并将插销头真正推进孔内。与前两个任务相比，"
        "其严格成功条件同时约束孔坐标系中的插入深度和横向误差，视觉上“碰到孔口”并不等于成功。",
    )
    add_para(
        doc,
        "实验首先复核 ManiSkill 3.0.1 的原生环境、43 维 state 观测、8 维动作、"
        "pd_joint_delta_pos 控制、100 步 horizon、dense reward 与严格 success predicate；随后在固定种子上"
        "建立随机策略和未训练官方 actor 基线。O1 完整复现官方 PPO-fast 75M 步配置，保留 92 个周期 checkpoint "
        "及 final，并通过全 checkpoint 筛选、独立 100 seeds 复验和 matched-seed 1,000 环境正式评估确定"
        " ckpt_2151 为 O1 最佳模型，严格成功率 17.5%。",
    )
    add_para(
        doc,
        "O1 的主要瓶颈不是抓取，而是最后插入深度不足。O2 从 O1 ckpt_2151 的不可变副本出发，"
        "重置优化器，以 1×10⁻⁴ 学习率线性退火继续训练 25M 步。O2 ckpt_551 在独立种子复验中达到"
        " 100/100，在 matched-seed 1,000 环境中达到 991/1,000；失败仅包括 1 次未抓取、1 次未完成"
        "预插入对准和 7 次深度不足。模型晋升后完成本地/服务器 SHA-256 一致性核对，并生成单视角、"
        "20 条展示视频及前 5 条四机位同步视频。",
    )
    add_para(doc, "关键词：ManiSkill；PPO；PegInsertionSide；机器人装配；插销插入；checkpoint 筛选；失败分类；多机位视频", indent=False)

    add_heading(doc, "一、任务目标与研究边界", 1)
    add_para(
        doc,
        "任务三的目标不是展示一次偶然成功，而是获得在严格原生判定下稳定、可复验、可交付的模型。"
        "训练、筛选和晋升遵循固定协议：训练前建立基线；保留每个 checkpoint；先用固定 100 seeds 全量筛选，"
        "再用独立 100 seeds 复验；最后在相同 1,000 环境随机实例上进行 matched-seed 比较。final 仅是时间上"
        "最后的 checkpoint，不自动视为 best。",
    )
    add_bullets(doc, [
        "服务器仅在 /vepfs-mlp2/queue013/public/huangborui 及其子目录写入；其他路径严格只读。",
        "PickCube 与 StackCube 的成果目录保持冻结；任务三使用独立 tasks、configs、runs、logs、videos 与报告命名空间。",
        "物理 GPU 0 当时存在活动计算进程，因此所有训练使用空闲的物理 GPU 1，单进程单卡执行。",
        "训练、评估和视频均以原生严格 success 为准，并保留失败阶段分类与负面实验。",
    ])

    add_heading(doc, "二、原生环境调查", 1)
    add_heading(doc, "2.1 观测、动作、控制与时限", 2)
    add_table(doc, ["项目", "复核结果", "说明"], [
        ("环境", "PegInsertionSide-v1", "ManiSkill 3.0.1 原生任务"),
        ("机器人", "panda_wristcam", "Panda 机械臂与腕部相机配置"),
        ("观测", "state，43 维", "训练使用低维状态，不使用图像观测"),
        ("动作", "8 维，[-1, 1]", "7 个机械臂关节控制量 + 1 个夹爪控制量"),
        ("控制模式", "pd_joint_delta_pos", "关节位置增量控制"),
        ("episode horizon", "100 步", "官方 PPO-fast 评估同样使用 100 步"),
    ], [1550, 2350, 5320])
    add_heading(doc, "2.2 严格成功判定", 2)
    add_para(
        doc,
        "原生环境将插销头位置变换到孔坐标系。记该位置为 (head_x, head_y, head_z)，"
        "孔的容许横向半径为 hole_radius = peg_radius + 0.003 m，则严格成功必须同时满足：",
    )
    add_callout(
        doc,
        "Native success",
        "head_x ≥ −0.015 m；|head_y| ≤ hole_radius；|head_z| ≤ hole_radius。"
        "三个条件必须在同一时刻成立。",
    )
    add_para(
        doc,
        "这一定义把“入口已对准”和“已经插入”明确区分开：即使插销轴线正确、孔口横向误差很小，"
        "若 head_x 未越过 −15 mm 阈值，仍严格判为失败。",
    )
    add_heading(doc, "2.3 Dense reward 的能力链", 2)
    add_para(
        doc,
        "官方 dense reward 依次鼓励末端接近插销尾部抓取位姿、在手指—物体夹角不超过约 20°时形成有效抓取、"
        "在目标坐标系中完成插销中心与插销头的 YZ 对准、推进插入，并在严格成功时把 raw dense reward 设为 10。"
        "normalized dense reward 再除以 10。奖励是连续训练信号，正式结论仍只采用原生 success。",
    )

    add_heading(doc, "三、官方 PPO-fast 参考与实验协议", 1)
    add_table(doc, ["参数", "官方/本项目值", "参数", "官方/本项目值"], [
        ("并行环境", "2,048", "rollout", "16"),
        ("update epochs", "8", "minibatches", "32"),
        ("gamma", "0.97", "GAE λ", "0.95"),
        ("总步数", "75M（O1）", "learning rate", "3×10⁻⁴（O1）"),
        ("评估步数", "100", "在线评估环境", "16"),
        ("设备", "物理 GPU 1", "拓扑", "单进程、单卡"),
    ], [1800, 2800, 1800, 2820])
    add_para(
        doc,
        "参考实现锁定官方 commit a4a4f9272ad64b1564035874b605ceb687b63ed8。"
        "O1 不改变官方任务和 PPO-fast 主参数。为了避免选择偏差，选择集、复验集和正式评估集使用不同 seed："
        "selection seed 0、independent seed 10000、formal matched seed 20260722。",
    )
    add_table(doc, ["阶段", "环境数", "seed", "用途"], [
        ("未训练基线", "100", "0", "确认任务难度与诊断链有效"),
        ("全 checkpoint 筛选", "每个 100", "0", "统一选择集排序"),
        ("独立复验", "每个 100", "10000", "排除对选择 seed 的偶然适配"),
        ("正式比较", "每个 1,000", "20260722", "旧候选与新候选 matched-seed 比较"),
    ], [2300, 1750, 1900, 3270])

    add_heading(doc, "四、训练前基线与诊断指标", 1)
    add_heading(doc, "4.1 未训练基线", 2)
    add_table(doc, ["策略", "严格成功", "抓取", "预插入对准", "入口对准", "平均回报"], [
        ("随机动作", "0/100", "0%", "0%", "0%", "2.089"),
        ("未训练官方 actor", "0/100", "0%", "0%", "0%", "2.593"),
    ], [2400, 1450, 1250, 1550, 1350, 1220])
    add_para(
        doc,
        "两组共 200 个 episode 均在“未抓取”阶段失败。这说明未经训练的 actor 不会因环境初始化或成功阈值"
        "宽松而偶然完成任务，也验证了失败分类能够识别最早断裂的能力阶段。",
    )
    add_heading(doc, "4.2 阶段化诊断指标", 2)
    add_table(doc, ["指标", "解释", "诊断目的"], [
        ("ever grasped", "episode 内是否曾形成有效抓取", "区分接近/抓取问题"),
        ("axis angle", "插销轴与目标轴的最小夹角", "衡量姿态对准"),
        ("pre-insert aligned", "入口上游的 YZ 预对准", "判断是否建立正确进孔路径"),
        ("entry aligned", "孔口处横向误差满足容差", "区分入口对准与插入深度"),
        ("maximum insertion x", "孔坐标系中最大带符号插入深度", "定位“差最后几毫米”"),
        ("strict success", "原生三条件同时满足", "唯一正式成功口径"),
    ], [2000, 3700, 3520])

    add_heading(doc, "五、O1：官方 75M 步从零训练", 1)
    add_heading(doc, "5.1 训练运行", 2)
    add_para(
        doc,
        "O1 run 为 runs/ppo-PegInsertionSide-v1-state-seed0-official75m_20260722_1630。"
        "训练从随机初始化开始，使用官方 75M 配置，在物理 GPU 1 上运行约 3 小时 27 分。"
        "共保留 92 个周期 checkpoint 和 final_ckpt.pt；训练过程中未自动删除旧模型，也未把 final 自动晋升。",
    )
    add_heading(doc, "5.2 全 checkpoint 筛选与独立复验", 2)
    add_table(doc, ["O1 checkpoint", "selection seed 0\n100 env", "independent seed 10000\n100 env", "说明"], [
        ("ckpt_2251", "26/100", "23/100", "选择集最高"),
        ("ckpt_2151", "21/100", "19/100", "正式评估后成为 O1 最佳"),
        ("ckpt_2126", "22/100", "14/100", "独立复验下降较大"),
        ("final", "11/100", "—", "明显低于多个中间 checkpoint"),
    ], [2100, 2100, 2500, 2520], highlight_rows={1: PALE_GREEN, 3: PALE_RED})
    add_callout(
        doc,
        "选择教训",
        "O1 final 在选择集仅 11/100，而多个周期 checkpoint 达到 21–26/100。"
        "这直接证明“最后保存的模型”不等于“泛化最好的模型”。",
        PALE_RED,
    )
    add_heading(doc, "5.3 Matched-seed 1,000 环境正式评估", 2)
    add_table(doc, ["O1 checkpoint", "严格成功", "成功率", "正式排序"], [
        ("ckpt_2151", "175/1,000", "17.5%", "1"),
        ("ckpt_2251", "160/1,000", "16.0%", "并列 2"),
        ("ckpt_2126", "160/1,000", "16.0%", "并列 2"),
        ("final", "89/1,000", "8.9%", "4"),
    ], [2600, 2200, 2000, 2420], highlight_rows={0: PALE_GREEN, 3: PALE_RED})
    add_para(
        doc,
        "虽然 ckpt_2251 在 100 环境选择集上领先，但 matched-seed 1,000 环境降低了抽样波动，"
        "ckpt_2151 最终以 17.5% 成为 O1 最佳。其模型 SHA-256 为 "
        "a350b940961c592d46729b7e9cc82bcedb8049cd1d114722d658c3e53afe2689。",
    )
    add_heading(doc, "5.4 O1 失败分析", 2)
    add_para(
        doc,
        "O1 已基本学会抓取与大方向对准，但大量轨迹在入口附近停止或回撤。阶段指标显示主要瓶颈已从"
        "“能否抓住”转移为“是否形成足够的正向插入深度”。因此，继续从零重复 75M 并非最有针对性的方案；"
        "更合理的做法是保留 O1 已获得的抓取、姿态和入口对准能力，以更小学习率继续优化最后插入阶段。",
    )

    add_heading(doc, "六、O2：从 O1 最佳候选继续训练", 1)
    add_heading(doc, "6.1 设计动机与训练参数", 2)
    add_para(
        doc,
        "O2 从 O1 ckpt_2151 的不可变副本启动，使用全新 run 与全新优化器状态，继续训练 25M 环境步。"
        "学习率设为 1×10⁻⁴，并线性退火至 0；其余环境规模、rollout、epochs、minibatches、gamma 和 GAE λ"
        "保持官方配置不变。该设计既保留 O1 已学会的子技能，又降低后期更新对成熟策略的破坏。",
    )
    add_table(doc, ["项目", "O1", "O2"], [
        ("初始化", "随机初始化", "O1 ckpt_2151 不可变副本"),
        ("优化器", "全新", "全新，不恢复旧动量"),
        ("训练步数", "75M", "追加 25M"),
        ("初始学习率", "3×10⁻⁴", "1×10⁻⁴"),
        ("学习率计划", "官方配置", "线性退火至 0"),
        ("运行时长", "约 3h27m", "约 1h50m"),
        ("保留模型", "92 periodic + final", "31 periodic + final + source copy"),
    ], [2600, 3310, 3310])
    add_para(
        doc,
        "O2 run 为 runs/ppo-PegInsertionSide-v1-state-seed0-official_ckpt2151_lr1e4_25m_20260723。"
        "训练末端在线评估达到 16/16，平均回报 85.62，但该在线结果仅作为健康指标，最终选择仍执行全 checkpoint 协议。",
    )

    add_heading(doc, "七、O2 全 checkpoint 选择与复验", 1)
    add_heading(doc, "7.1 固定 selection seed 0", 2)
    add_table(doc, ["O2 checkpoint", "100 环境严格成功", "备注"], [
        ("ckpt_551", "99/100", "候选"),
        ("ckpt_626", "99/100", "候选"),
        ("ckpt_701", "99/100", "候选"),
        ("ckpt_726", "99/100", "候选"),
        ("ckpt_751", "99/100", "候选"),
        ("final", "97/100", "不自动晋升"),
    ], [3000, 2900, 3320], highlight_rows={0: PALE_GREEN, 5: PALE_RED})
    add_heading(doc, "7.2 独立 seed 10000 复验", 2)
    add_table(doc, ["O2 checkpoint", "100 环境严格成功", "独立复验判断"], [
        ("ckpt_551", "100/100", "进入正式评估"),
        ("ckpt_701", "100/100", "进入正式评估"),
        ("ckpt_626", "99/100", "高稳定候选"),
        ("ckpt_751", "99/100", "高稳定候选"),
        ("ckpt_726", "98/100", "略低"),
    ], [3000, 2900, 3320], highlight_rows={0: PALE_GREEN, 1: PALE_GREEN})
    add_para(
        doc,
        "选择集上五个 checkpoint 同为 99/100，无法可靠区分。独立复验把 ckpt_551 与 ckpt_701 "
        "同时确认到 100/100，因此二者进入 matched-seed 1,000 环境正式比较。该流程避免了依据单个 100 环境"
        "样本中 1–2 次随机失败做过度判断。",
    )

    add_heading(doc, "八、正式 1,000 环境结果与 O1→O2 提升", 1)
    add_heading(doc, "8.1 严格成功率", 2)
    add_table(doc, ["模型", "严格成功", "成功率", "相对 O1 最佳"], [
        ("O2 ckpt_551", "991/1,000", "99.1%", "+81.6 pp"),
        ("O2 ckpt_701", "990/1,000", "99.0%", "+81.5 pp"),
        ("O1 ckpt_2151", "175/1,000", "17.5%", "基线"),
        ("O1 final", "89/1,000", "8.9%", "−8.6 pp"),
    ], [2800, 2000, 1800, 2620], highlight_rows={0: PALE_GREEN, 2: PALE_BLUE, 3: PALE_RED})
    add_callout(
        doc,
        "核心结果",
        "O2 ckpt_551 在与 O1 候选完全 matched 的 1,000 个环境实例上达到 99.1%，"
        "相较 O1 最佳 17.5% 绝对提升 81.6 个百分点。",
        PALE_GREEN,
    )
    add_heading(doc, "8.2 O2 ckpt_551 阶段指标", 2)
    add_table(doc, ["指标", "结果", "解读"], [
        ("严格成功", "991/1,000（99.1%）", "正式 native success"),
        ("ever grasped", "99.9%", "抓取已基本饱和"),
        ("姿态 ≤10°", "99.9%", "轴向对准稳定"),
        ("pre-insert aligned", "99.8%", "绝大多数建立正确进孔路径"),
        ("entry aligned", "100.0%", "全部轨迹至少曾满足孔口横向对准"),
        ("最小轴角均值", "0.836°", "姿态精度高"),
        ("最小横向误差均值", "1.306 mm", "入口对准达到毫米级"),
        ("最大 insertion x 均值", "+26.46 mm", "整体具备明显正插入深度"),
        ("normalized-dense return", "83.154", "与高成功率一致"),
    ], [2600, 2800, 3820])

    add_heading(doc, "九、互斥失败分类与评估修正", 1)
    add_heading(doc, "9.1 O2 ckpt_551 的 9 个失败", 2)
    add_table(doc, ["互斥失败类别", "数量", "占全部 1,000", "含义"], [
        ("未抓取", "1", "0.1%", "未建立有效抓取"),
        ("未完成预插入对准", "1", "0.1%", "抓取后未形成上游对准"),
        ("插入深度不足", "7", "0.7%", "入口已对准但未跨越严格 x 阈值"),
        ("其他类别", "0", "0.0%", "无未解释失败"),
    ], [3100, 1300, 1900, 2920], highlight_rows={2: PALE_RED})
    add_para(
        doc,
        "剩余误差高度集中：9 个失败中 7 个仍是最后插入深度不足。与 O1 相比，O2 并非仅把回报整体提高，"
        "而是把抓取、姿态、预对准和入口对准依次推到接近饱和，最终仅留下极少量深度边界失败。",
    )
    add_heading(doc, "9.2 失败掩码重叠问题及修正", 2)
    add_para(
        doc,
        "首次正式评估后发现旧诊断脚本的失败 mask 存在重叠：同一 episode 可能被多个后续类别同时计数。"
        "该问题不影响原生 strict success，也不改变 checkpoint 排名，但会使失败原因总数失真。"
        "脚本随后修正为按最早断裂阶段互斥分类：未抓取、未姿态对准、未预插入对准、未入口对准、"
        "插入深度不足、其他。",
    )
    add_bullets(doc, [
        "服务器原评估脚本先备份到 backups/peg_insertion_side_eval_before_mutually_exclusive_fix_20260723。",
        "ckpt_551 与 ckpt_701 的 1,000 环境正式评估均写入全新的 mutually_exclusive_20260723 目录。",
        "重跑后 strict success 与排序完全不变；修正只提高失败解释的统计一致性。",
    ])

    # Keep this self-contained explanatory section together instead of leaving
    # its concluding callout orphaned at the top of the following page.
    doc.add_page_break()
    add_heading(doc, "十、为什么 O2 会出现非线性跃升", 1)
    add_para(
        doc,
        "O2 从 17.5% 跃升到 99.1% 看似突然，但阶段指标表明它并不是“凭空变强”。O1 已完成最困难的"
        "探索性子技能：靠近、抓取、抬起、改变插销姿态并到达孔口附近。严格 success 是阈值型指标；当大量"
        "轨迹聚集在阈值外侧时，策略质量的小幅连续改善会使成功率产生很大的离散跃迁。",
    )
    add_table(doc, ["机制", "O1 状态", "O2 的作用"], [
        ("技能继承", "已有抓取与入口附近行为", "避免从零重新探索完整能力链"),
        ("更低学习率", "后期更新仍可能破坏成熟子技能", "1e-4 + 退火让动作分布逐步精修"),
        ("新优化器", "不继承旧动量与优化历史", "在保留权重的同时重新建立更新尺度"),
        ("额外 25M 数据", "深度不足样本仍多", "反复学习对准后持续正向推进"),
        ("阈值效应", "大量轨迹停在成功边界附近", "几毫米改善转化为大批严格成功"),
    ], [2100, 3400, 3720])
    add_callout(
        doc,
        "解释边界",
        "不能把提升简单归因于“多训练 25M 步”。有效组合是：从正式评估选出的 O1 ckpt_2151 "
        "开始、重置优化器、降低学习率并退火、继续保留全 checkpoint 后重新筛选。",
    )

    add_heading(doc, "十一、模型晋升、哈希与可复现交付", 1)
    add_heading(doc, "11.1 晋升规则与最终模型", 2)
    add_para(
        doc,
        "仅在 ckpt_551 经全 checkpoint 选择、独立 100 seeds 复验并在 matched-seed 1,000 环境中"
        "确认优于旧候选后，才复制晋升为 best。源 checkpoint、其他周期 checkpoint 与 final 均继续保留。",
    )
    add_table(doc, ["位置", "路径"], [
        ("本地", "checkpoints/ppo_peg_insertion_side_official_ckpt2151_lr1e4_25m/\nbest_o2_ckpt551_991of1000_matchedseed.pt"),
        ("服务器", "runs/ppo-PegInsertionSide-v1-state-seed0-official_ckpt2151_lr1e4_25m_20260723/\nbest_o2_ckpt551_991of1000_matchedseed.pt"),
    ], [1800, 7420], font_size=8.5)
    add_callout(
        doc,
        "SHA-256",
        "b48a4e0732de5e1e68bc906897166be95f2cecb3a4c8d245680dc62e2e7a6c49。"
        "本地与服务器已核对一致。",
        PALE_GREEN,
    )
    add_heading(doc, "11.2 本地测试环境", 2)
    add_para(
        doc,
        "本地独立环境为 .venv-peg-insertion-side，包含 ManiSkill 3.0.1、SAPIEN、PyTorch 与视频依赖。"
        "Apple Silicon 使用 CPU PhysX；测试脚本为 scripts/test_peg_insertion_side_local.py，支持指定 episode、"
        "seed、视频路径和 hero/front/side/top 机位。",
    )
    add_callout(
        doc,
        "快速测试命令",
        ".venv-peg-insertion-side/bin/python scripts/test_peg_insertion_side_local.py "
        "--episodes 10 --seed 0",
        PALE_BLUE,
    )
    add_callout(
        doc,
        "单条录像命令",
        ".venv-peg-insertion-side/bin/python scripts/test_peg_insertion_side_local.py "
        "--episodes 1 --seed 1 --video videos/peg_insertion_side/local_test/seed1.mp4",
        PALE_BLUE,
    )

    add_heading(doc, "十二、视频证据与多机位展示", 1)
    add_heading(doc, "12.1 单条与 20 条成功展示", 2)
    add_table(doc, ["视频集合", "数量", "编码与尺寸", "验证"], [
        ("最佳模型 seed0", "1", "H.264 · 512×512 · 30 fps", "101 帧 · 3.3667 s · strict 1/1"),
        ("showcase seeds 0–19", "20", "H.264 · 512×512 · 30 fps", "每条 101 帧 · 全部严格成功"),
    ], [2600, 1200, 2700, 2720])
    add_para(
        doc,
        "20 条展示视频位于 videos/peg_insertion_side/showcase_best_o2_ckpt551_20_successes_20260723/"
        "successes/，并配套 ffprobe_manifest.csv。所有视频均用 ffprobe 核对编码、分辨率、帧率、帧数与时长。",
    )
    add_heading(doc, "12.2 前五条同步四机位", 2)
    add_para(
        doc,
        "对 seeds 0–4 的同一轨迹分别生成 hero、front、side、top 四个独立视角，并额外生成 2×2 四宫格。"
        "四宫格布局为左上 hero、右上 front、左下 side、右下 top。",
    )
    add_table(doc, ["产物", "数量", "规格"], [
        ("独立机位视频", "20", "4 机位 × 5 seeds；512×512；H.264；30 fps；101 帧"),
        ("四宫格视频", "5", "1024×1024；H.264；30 fps；101 帧"),
        ("严格成功日志", "20", "每个 seed/机位均保留运行日志"),
        ("ffprobe manifest", "1", "统一记录所有视频媒体参数"),
    ], [2900, 1400, 4920])
    add_para(
        doc,
        "本地目录为 videos/peg_insertion_side/multiview_best_o2_ckpt551_first5_20260723/，"
        "服务器存在同名任务三视频命名空间副本。多机位不是四次不同 rollout，而是用于观察同一确定性策略轨迹"
        "在不同视角下的抓取、姿态调整、孔口对准与插入过程。",
    )

    add_heading(doc, "十三、负面结果、风险与研究记录", 1)
    add_table(doc, ["记录项", "结果", "保留价值"], [
        ("随机策略基线", "0/100，全部未抓取", "证明环境不会轻易偶然成功"),
        ("未训练 actor", "0/100，全部未抓取", "验证初始化不具备任务技能"),
        ("O1 final", "89/1,000", "证明 final 不等于 best"),
        ("O1 ckpt_2251", "选择集 26%，正式仅 16%", "证明小样本排序可能反转"),
        ("O1 主要失败", "入口附近深度不足", "为 O2 continuation 提供针对性依据"),
        ("旧失败 mask", "类别可能重叠", "修正统计口径并重跑，不掩盖诊断缺陷"),
        ("O2 剩余失败", "7/9 为深度不足", "指出 99.1% 后仍存在的边界风险"),
    ], [2450, 2900, 3870])
    add_para(
        doc,
        "当前结论适用于 ManiSkill 3.0.1、指定机器人、state 观测、pd_joint_delta_pos 控制与当前随机化分布。"
        "99.1% 是 matched-seed 1,000 环境中的仿真成功率，不应直接外推为真实机器人成功率。若后续开展 sim-to-real，"
        "还需补充视觉观测、动力学与接触随机化、控制延迟、标定误差和真实夹持摩擦等验证。",
    )

    add_heading(doc, "十四、结论", 1)
    add_para(
        doc,
        "任务三已经形成一条完整且可审计的研究闭环：原生任务复核 → 固定种子未训练基线 → 官方 75M 步 O1 "
        "从零训练 → 全 checkpoint 筛选 → 独立种子复验 → 1,000 环境正式比较 → O1 失败阶段定位 → "
        "低学习率 O2 continuation → 再次全量筛选与正式评估 → 模型晋升、哈希核对和视频交付。",
    )
    add_para(
        doc,
        "最终 O2 ckpt_551 达到 991/1,000（99.1%）严格成功率，相较 O1 最佳 175/1,000 提升"
        " 81.6 个百分点。其抓取、姿态对准、预插入对准与入口对准均接近饱和，剩余失败主要集中在极少数"
        "插入深度不足轨迹。结果同时说明：复杂机械装配任务的关键不仅是训练更久，更是采用严格成功谓词、"
        "阶段化诊断、保留并筛选全部 checkpoint，以及用独立与 matched-seed 评估把真实提升和抽样波动分开。",
    )
    add_callout(
        doc,
        "最终交付状态",
        "最佳模型、评估记录、失败分类、单视角与多机位视频、本地测试环境和任务报告均已形成独立任务三命名空间；"
        "PickCube 与 StackCube 成果目录保持冻结。",
        PALE_GREEN,
    )

    add_heading(doc, "附录 A：关键路径索引", 1)
    add_table(doc, ["类别", "本地路径"], [
        ("任务说明", "tasks/peg_insertion_side/README.md"),
        ("官方配置", "configs/peg_insertion_side/official_ppo_fast_75m.yaml"),
        ("实验日志", "docs/peg_insertion_side_experiment_log.md"),
        ("正式评估脚本", "scripts/evaluate_official_peg_insertion_side.py"),
        ("本地测试脚本", "scripts/test_peg_insertion_side_local.py"),
        ("最佳模型", "checkpoints/ppo_peg_insertion_side_official_ckpt2151_lr1e4_25m/\nbest_o2_ckpt551_991of1000_matchedseed.pt"),
        ("20 条展示", "videos/peg_insertion_side/showcase_best_o2_ckpt551_20_successes_20260723/"),
        ("四机位展示", "videos/peg_insertion_side/multiview_best_o2_ckpt551_first5_20260723/"),
    ], [2100, 7120], font_size=8.5)

    add_heading(doc, "附录 B：正式评估数据摘要", 1)
    add_table(doc, ["模型", "selection 100", "independent 100", "formal 1,000", "结论"], [
        ("O1 final", "11", "—", "89", "不晋升"),
        ("O1 ckpt_2251", "26", "23", "160", "选择集领先但正式非最佳"),
        ("O1 ckpt_2151", "21", "19", "175", "O1 最佳、O2 起点"),
        ("O2 final", "97", "—", "—", "不自动晋升"),
        ("O2 ckpt_701", "99", "100", "990", "正式亚军"),
        ("O2 ckpt_551", "99", "100", "991", "最终晋升"),
    ], [2300, 1650, 1900, 1800, 1570], highlight_rows={2: PALE_BLUE, 5: PALE_GREEN})
    add_para(
        doc,
        "注：selection 与 independent 列的数值均为 100 环境中的成功次数；formal 列为同一 matched seed "
        "下 1,000 环境中的成功次数。破折号表示该项未作为正式选择依据，而非推定为零。",
        indent=False,
    )

    # Prevent accidental overwrite if this builder is re-run.
    if OUTPUT.exists():
        raise FileExistsError(f"Refusing to overwrite existing report: {OUTPUT}")
    doc.core_properties.title = "PegInsertionSide-v1 任务三完整研究报告"
    doc.core_properties.subject = "ManiSkill 3.0.1 PPO-fast task-3 final report"
    doc.core_properties.author = "黄泊睿"
    doc.core_properties.keywords = "ManiSkill, PPO, PegInsertionSide, robot insertion"
    doc.settings.element.append(deepcopy(OxmlElement("w:updateFields")))
    doc.settings.element[-1].set(qn("w:val"), "true")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
